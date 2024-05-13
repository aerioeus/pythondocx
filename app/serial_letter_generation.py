
import os
import pandas as pd
from docx import Document
import logging
from docx2pdf import convert
import win32com.client


template_path = 'template.docx'
excel_path = 'data.xlsx'
docx_folder = 'letters'
pdf_output_folder = 'pdfs'

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Change current working directory to the script's directory
os.chdir(os.path.dirname(os.path.abspath(__file__)))

# Function to load Excel data into a DataFrame
def load_excel_data(file_path: str, sheet_name: str = 'Sheet1') -> pd.DataFrame:
    return pd.read_excel(file_path, sheet_name=sheet_name)

# Function to create a personalized letter from the template
def create_personalized_letter(template_path: str, data: dict, output_path: str) -> None:
    # Verify the template file exists
    if not os.path.exists(template_path):
        logging.error(f"Template file '{template_path}' not found.")
        raise FileNotFoundError(f"Template file '{template_path}' not found.")

    try:
        # Load the document template
        doc = Document(template_path)
    except Exception as e:
        logging.error(f"Failed to load document template from {template_path}: {e}")
        raise

    try:
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, replacement in data.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(replacement))
                    logging.debug(f"Replaced placeholder {placeholder} in paragraph.")

        # Replace placeholders in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, replacement in data.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(replacement))
                            logging.debug(f"Replaced placeholder {placeholder} in table cell.")

        # Save the modified document
        doc.save(output_path)
        logging.info(f"Document saved successfully at {output_path}")
    except Exception as e:
        logging.error(f"An error occurred while creating or saving the document: {e}")
        raise

# Function to count files in a directory by extension
def count_files_in_directory(directory: str, file_extension: str) -> int:
    try:
        files = os.listdir(directory)
        return sum(1 for file in files if file.endswith(file_extension))
    except Exception as e:
        logging.error(f"Error counting files with extension {file_extension} in {directory}: {e}")
        raise  # Re-raise the exception to be caught by the calling function if needed

# Function to generate serial letters and count the generated DOCX files
def generate_serial_letters_with_count(template_path: str, excel_path: str, output_dir: str, sheet_name: str = 'Sheet1') -> None:
    # Verify existence of the Excel file
    if not os.path.exists(excel_path):
        logging.error(f"Excel file '{excel_path}' not found.")
        raise FileNotFoundError(f"Excel file '{excel_path}' not found.")

    # Load data from Excel
    try:
        df = load_excel_data(excel_path, sheet_name)
    except Exception as e:
        logging.error(f"Failed to load data from {excel_path}: {e}")
        return  # Exit the function if data cannot be loaded

    # Ensure the output directory exists
    try:
        os.makedirs(output_dir, exist_ok=True)
    except Exception as e:
        logging.error(f"Failed to create or access output directory {output_dir}: {e}")
        return  # Exit the function if the directory cannot be created or accessed

    # Generate personalized letters
    for index, row in df.iterrows():
        try:
            data = {f"{{{col}}}": val for col, val in row.items()}
            output_path = os.path.join(output_dir, f"{row['Firstname']}_letter.docx")
            create_personalized_letter(template_path, data, output_path)
            logging.info(f"Created: {output_path}")
        except Exception as e:
            logging.error(f"Failed to create letter for {row['Firstname']}: {e}")

    # Count the generated DOCX files
    try:
        docx_count = count_files_in_directory(output_dir, '.docx')
        logging.info(f"Total DOCX files created: {docx_count}")
    except Exception as e:
        logging.error(f"Failed to count DOCX files in {output_dir}: {e}")

    logging.info("Serial letter generation and counting complete.")

# Function to convert all .docx files to PDFs and count the PDF files
def convert_all_word_files_to_pdfs_with_count(word_dir: str, pdf_dir: str) -> None:
    # Ensure the PDF output directory exists
    try:
        if not os.path.exists(pdf_dir):
            os.makedirs(pdf_dir)
    except Exception as e:
        logging.error(f"Failed to create directory {pdf_dir}: {e}")
        return  # Exit the function if the directory cannot be created

    def close_word():
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Quit()
        except Exception as e:
            logging.error("Failed to close Word: {}".format(e))

    # Try to convert all .docx files to PDFs
    try:
        convert(word_dir, pdf_dir)
    except Exception as e:
        logging.error(f"Failed to convert files from {word_dir} to {pdf_dir}: {e}")
    finally:
        close_word()  # Ensure Word is closed after conversion attempt

    # Count the PDF files and log the count
    try:
        pdf_count = count_files_in_directory(pdf_dir, '.pdf')
        logging.info(f"Total PDF files created: {pdf_count}")
    except Exception as e:
        logging.error(f"Failed to count PDF files in {pdf_dir}: {e}")

    logging.info("Conversion complete.")


if not os.path.exists(template_path):
    raise FileNotFoundError(f"Template file '{template_path}' not found.")

if not os.path.exists(excel_path):
    raise FileNotFoundError(f"Excel file '{excel_path}' not found.")

generate_serial_letters_with_count(template_path, excel_path, docx_folder)
convert_all_word_files_to_pdfs_with_count(docx_folder, pdf_output_folder)
