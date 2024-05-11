
import os
import pandas as pd
from docx import Document
from docx2pdf import convert
import win32com.client


template_path = 'template.docx'
excel_path = 'data.xlsx'
docx_folder = 'letters'
pdf_output_folder = 'pdfs'


# Change current working directory to the script's directory
os.chdir(os.path.dirname(os.path.abspath(__file__)))

# Function to load Excel data into a DataFrame
def load_excel_data(file_path: str, sheet_name: str = 'Sheet1') -> pd.DataFrame:
    return pd.read_excel(file_path, sheet_name=sheet_name)

# Function to create a personalized letter from the template
def create_personalized_letter(template_path: str, data: dict, output_path: str) -> None:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file '{template_path}' not found.")

    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for placeholder, replacement in data.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(replacement))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, replacement in data.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(replacement))

    doc.save(output_path)

# Function to count files in a directory by extension
def count_files_in_directory(directory: str, file_extension: str) -> int:
    files = os.listdir(directory)
    return sum(1 for file in files if file.endswith(file_extension))

# Function to generate serial letters and count the generated DOCX files
def generate_serial_letters_with_count(template_path: str, excel_path: str, output_dir: str, sheet_name: str = 'Sheet1') -> None:
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Excel file '{excel_path}' not found.")

    df = load_excel_data(excel_path, sheet_name)
    os.makedirs(output_dir, exist_ok=True)

    for _, row in df.iterrows():
        data = {f"{{{col}}}": val for col, val in row.items()}
        output_path = os.path.join(output_dir, f"{row['Name']}_letter.docx")
        create_personalized_letter(template_path, data, output_path)
        print(f"Created: {output_path}")

    docx_count = count_files_in_directory(output_dir, '.docx')
    print(f"Total DOCX files created: {docx_count}")

# Function to convert all .docx files to PDFs and count the PDF files
def convert_all_word_files_to_pdfs_with_count(word_dir: str, pdf_dir: str) -> None:
    if not os.path.exists(pdf_dir):
        os.makedirs(pdf_dir)

    def close_word():
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Quit()
        except Exception as e:
            print("Failed to close Word:", e)

    try:
        convert(word_dir, pdf_dir)
    finally:
        close_word()

    pdf_count = count_files_in_directory(pdf_dir, '.pdf')
    print(f"Total PDF files created: {pdf_count}")
    print("Conversion complete.")


if not os.path.exists(template_path):
    raise FileNotFoundError(f"Template file '{template_path}' not found.")

if not os.path.exists(excel_path):
    raise FileNotFoundError(f"Excel file '{excel_path}' not found.")

generate_serial_letters_with_count(template_path, excel_path, docx_folder)
convert_all_word_files_to_pdfs_with_count(docx_folder, pdf_output_folder)
